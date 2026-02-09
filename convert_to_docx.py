from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_table_borders(table):
    """Add borders to table"""
    tbl = table._tbl
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tbl.tblPr.append(tblBorders)

def create_ulp_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('Dokumen Konseptual: Unified Learning Platform (ULP)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Arsitektur dan Tata Kelola', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    meta = doc.add_paragraph()
    meta.add_run('Status Dokumen: ').bold = True
    meta.add_run('Draft untuk Diskusi Internal\n')
    meta.add_run('Tanggal: ').bold = True
    meta.add_run('Februari 2026\n')
    meta.add_run('Versi: ').bold = True
    meta.add_run('0.2 — Disesuaikan dengan Implementasi Pilot')
    
    doc.add_paragraph()
    
    # 1. Pendahuluan
    doc.add_heading('1. Pendahuluan', level=1)
    
    doc.add_heading('1.1 Latar Belakang', level=2)
    doc.add_paragraph(
        'Organisasi saat ini mengoperasikan beberapa sistem pembelajaran yang berkembang secara organik '
        'sesuai kebutuhan masing-masing unit. Kondisi ini menghasilkan fragmentasi yang dapat menyulitkan '
        'pegawai dalam menemukan dan mengakses materi pembelajaran yang relevan.'
    )
    doc.add_paragraph(
        'Inisiatif Unified Learning Platform (ULP) hadir sebagai upaya untuk menyediakan satu pintu masuk terpadu '
        'bagi pegawai dalam mengakses berbagai sumber pembelajaran. Platform ini dirancang untuk menggantikan '
        'sistem LXP (pintar.setneg.go.id) dan LMS (belajar.setneg.go.id) yang saat ini berjalan terpisah.'
    )
    
    doc.add_heading('1.2 Tujuan Dokumen', level=2)
    doc.add_paragraph('Dokumen ini disusun untuk:')
    doc.add_paragraph('• Menyamakan persepsi antara unit terkait, inisiator, dan Badan Teknologi Data dan Informasi terkait konsep ULP', style='List Bullet')
    doc.add_paragraph('• Menjelaskan fitur-fitur yang sudah diimplementasikan dalam pilot', style='List Bullet')
    doc.add_paragraph('• Membuka ruang diskusi konstruktif sebelum pengambilan keputusan teknis lanjutan', style='List Bullet')
    
    doc.add_heading('1.3 Penegasan Penting', level=2)
    warning = doc.add_paragraph()
    warning.add_run('Dokumen ini bukan desain final. ').bold = True
    warning.add_run('Seluruh isi masih memerlukan validasi, penyesuaian, serta persetujuan dari pemangku kepentingan—terutama Badan Teknologi Data dan Informasi.')
    
    # 2. Ruang Lingkup
    doc.add_heading('2. Ruang Lingkup & Batasan', level=1)
    
    doc.add_heading('2.1 Yang Dicakup oleh ULP (Implementasi Saat Ini)', level=2)
    
    table1 = doc.add_table(rows=8, cols=2)
    add_table_borders(table1)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    hdr = table1.rows[0].cells
    hdr[0].text = 'Modul'
    hdr[1].text = 'Deskripsi'
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True
    
    data1 = [
        ('Kursus & Modul', 'Struktur pembelajaran dengan video, transkrip, dan materi pendukung'),
        ('Quiz & Asesmen', 'Pre-test, post-test, dan quiz dengan berbagai tipe soal'),
        ('YouTube Integration', 'Import playlist YouTube sebagai konten pembelajaran'),
        ('Learning Analytics (xAPI)', 'Pelacakan aktivitas belajar untuk keperluan pelaporan'),
        ('PBGM', 'Pembelajaran Berbasis Gagasan dan Masalah (project-based learning)'),
        ('Forum Diskusi', 'Ruang diskusi per kursus'),
        ('Sesi & Kehadiran', 'Manajemen sesi tatap muka dan pelacakan kehadiran'),
    ]
    
    for i, (modul, desc) in enumerate(data1):
        row = table1.rows[i + 1].cells
        row[0].text = modul
        row[1].text = desc
    
    doc.add_paragraph()
    
    doc.add_heading('2.2 Yang Secara Eksplisit TIDAK Dilakukan oleh ULP', level=2)
    
    table2 = doc.add_table(rows=5, cols=2)
    add_table_borders(table2)
    
    hdr2 = table2.rows[0].cells
    hdr2[0].text = 'Aspek'
    hdr2[1].text = 'Posisi ULP'
    for cell in hdr2:
        cell.paragraphs[0].runs[0].bold = True
    
    data2 = [
        ('Pengelolaan identitas (SSO)', 'Menggunakan autentikasi existing (NextAuth + LDAP organisasi)'),
        ('Penyimpanan data master pegawai', 'Mengambil dari sistem HR yang ada'),
        ('Penilaian resmi & keputusan kepegawaian', 'Tidak terlibat — tetap di sistem existing'),
        ('Penerbitan sertifikat resmi', 'Sertifikat ULP bersifat internal sebagai bukti penyelesaian'),
    ]
    
    for i, (aspek, posisi) in enumerate(data2):
        row = table2.rows[i + 1].cells
        row[0].text = aspek
        row[1].text = posisi
    
    doc.add_paragraph()
    
    # 3. Arsitektur Konseptual
    doc.add_heading('3. Arsitektur Konseptual', level=1)
    
    doc.add_heading('3.1 Model Hubungan Sistem', level=2)
    
    # Diagram as formatted text
    diagram = doc.add_paragraph()
    diagram.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    diagram_text = """
                    ┌─────────────────┐
                    │    PENGGUNA     │
                    │  (Learner/      │
                    │   Instructor)   │
                    └────────┬────────┘
                             │
               ┌─────────────▼─────────────┐
               │ UNIFIED LEARNING PLATFORM │
               │          (ULP)            │
               │  ┌───────────────────┐    │
               │  │Kursus│Quiz │PBGM  │    │
               │  │Modul │Forum│Sesi  │    │
               │  │xAPI Analytics     │    │
               │  └───────────────────┘    │
               └──┬────────┬────────┬──────┘
                  │        │        │
    ┌─────────────┼────────┼────────┘
    ▼             ▼        ▼
┌────────┐  ┌──────────┐  ┌───────────────┐
│YouTube │  │LDAP/     │  │   Database    │
│(Konten)│  │Google    │  │  Kepegawaian  │
└────────┘  └──────────┘  └───────────────┘
    │            │               │
  Read         Read            Read
  Only         Only            Only
"""
    run = diagram.add_run(diagram_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    
    doc.add_heading('3.2 Prinsip Arsitektur yang Diterapkan', level=2)
    
    doc.add_paragraph('1. Human-in-Control', style='List Number')
    doc.add_paragraph('Semua keputusan substantif oleh manusia. AI hanya membantu drafting konten.')
    
    doc.add_paragraph('2. Separation of Concern', style='List Number')
    doc.add_paragraph('• Autentikasi: NextAuth + LDAP + Google OAuth (bukan ULP)', style='List Bullet')
    doc.add_paragraph('• Konten: YouTube import dan konten manual (upload langsung)', style='List Bullet')
    doc.add_paragraph('• Tracking: xAPI untuk audit trail', style='List Bullet')
    
    doc.add_paragraph('3. Badan Teknologi Data dan Informasi sebagai Pemilik System-of-Record', style='List Number')
    doc.add_paragraph('ULP tidak menyimpan data master. Data pegawai tetap di Database Kepegawaian.')
    
    doc.add_paragraph('4. Reversibility', style='List Number')
    doc.add_paragraph('Komponen dapat dimatikan tanpa mengganggu sistem lain.')
    
    doc.add_paragraph('5. Incremental Adoption', style='List Number')
    doc.add_paragraph('Dimulai dari pilot terbatas. Fitur ditambahkan bertahap sesuai kebutuhan.')
    
    doc.add_heading('3.3 Posisi AI dalam ULP', level=2)
    doc.add_paragraph('AI diposisikan sebagai lapisan pendukung non-kritikal:')
    
    table3 = doc.add_table(rows=6, cols=2)
    add_table_borders(table3)
    
    hdr3 = table3.rows[0].cells
    hdr3[0].text = '✅ Penggunaan yang Diimplementasikan'
    hdr3[1].text = '❌ TIDAK Digunakan Untuk'
    for cell in hdr3:
        cell.paragraphs[0].runs[0].bold = True
    
    data3 = [
        ('Generate quiz dari transkrip video', 'Autentikasi pengguna'),
        ('Refine judul lesson untuk kejelasan', 'Otorisasi akses'),
        ('Generate deskripsi kursus', 'Penilaian resmi final'),
        ('Generate thumbnail', 'Keputusan kepegawaian'),
        ('Learning analytics insights', 'Penerbitan sertifikat resmi'),
    ]
    
    for i, (ya, tidak) in enumerate(data3):
        row = table3.rows[i + 1].cells
        row[0].text = ya
        row[1].text = tidak
    
    doc.add_paragraph()
    doc.add_paragraph('Catatan Implementasi:', style='Intense Quote')
    doc.add_paragraph('• AI menggunakan proxy service (AI_PROXY_URL)', style='List Bullet')
    doc.add_paragraph('• Output AI selalu dapat diedit oleh instructor sebelum dipublikasikan', style='List Bullet')
    doc.add_paragraph('• Quiz AI-generated ditandai dengan flag isAIGenerated', style='List Bullet')
    doc.add_paragraph('• Tidak ada automated decision-making untuk penilaian atau pass/fail', style='List Bullet')
    
    # 4. Tata Kelola
    doc.add_heading('4. Tata Kelola & Ownership', level=1)
    
    doc.add_heading('4.1 Pembagian Peran', level=2)
    
    table4 = doc.add_table(rows=4, cols=2)
    add_table_borders(table4)
    
    hdr4 = table4.rows[0].cells
    hdr4[0].text = 'Pihak'
    hdr4[1].text = 'Tanggung Jawab'
    for cell in hdr4:
        cell.paragraphs[0].runs[0].bold = True
    
    data4 = [
        ('Inisiator', 'Menyusun konsep, mengoperasikan pilot, dokumentasi'),
        ('PPKASN, Pusat Pembinaan JF Analis Kerja Sama, dan Pusat Pembinaan JF Penerjemah', 'Input kebutuhan kurikulum, validasi konten'),
        ('Badan Teknologi Data dan Informasi', 'Review integrasi LDAP, standar keamanan, infrastruktur'),
    ]
    
    for i, (pihak, tanggung) in enumerate(data4):
        row = table4.rows[i + 1].cells
        row[0].text = pihak
        row[1].text = tanggung
    
    doc.add_paragraph()
    
    doc.add_heading('4.2 Mekanisme Pengambilan Keputusan', level=2)
    doc.add_paragraph('• Keputusan arsitektur teknis: Badan Teknologi Data dan Informasi', style='List Bullet')
    doc.add_paragraph('• Keputusan konten dan kurikulum: PPKASN, Pusat Pembinaan JF Analis Kerja Sama, dan Pusat Pembinaan JF Penerjemah', style='List Bullet')
    doc.add_paragraph('• Keputusan lanjut/hentikan pilot: Forum bersama', style='List Bullet')
    
    doc.add_heading('4.3 Posisi Pilot', level=2)
    doc.add_paragraph('Tahap pilot saat ini:')
    doc.add_paragraph('• Lingkup pengguna terbatas', style='List Bullet')
    doc.add_paragraph('• Tidak ada SLA formal', style='List Bullet')
    doc.add_paragraph('• Sertifikat bersifat internal (bukan pengganti sertifikat resmi)', style='List Bullet')
    doc.add_paragraph('• Dapat dihentikan kapan saja tanpa dampak operasional sistem existing', style='List Bullet')
    
    # 5. Manajemen Risiko
    doc.add_heading('5. Manajemen Risiko & Kepatuhan', level=1)
    
    doc.add_heading('5.1 Mitigasi Risiko', level=2)
    
    table5 = doc.add_table(rows=5, cols=2)
    add_table_borders(table5)
    
    hdr5 = table5.rows[0].cells
    hdr5[0].text = 'Aspek'
    hdr5[1].text = 'Implementasi'
    for cell in hdr5:
        cell.paragraphs[0].runs[0].bold = True
    
    data5 = [
        ('Autentikasi', 'Menggunakan LDAP organisasi, bukan sistem terpisah'),
        ('Data Privacy', 'Video dari YouTube (publik), data tracking internal'),
        ('Dependency', 'YouTube sebagai read-only source, dapat dialihkan ke sumber lain'),
        ('AI Output', 'Selalu dapat diedit manusia, ditandai sebagai AI-generated'),
    ]
    
    for i, (aspek, imp) in enumerate(data5):
        row = table5.rows[i + 1].cells
        row[0].text = aspek
        row[1].text = imp
    
    doc.add_paragraph()
    
    doc.add_heading('5.2 Audit Trail', level=2)
    doc.add_paragraph('• Setiap aktivitas belajar dicatat via xAPI statements', style='List Bullet')
    doc.add_paragraph('• Log mencakup: user, verb, object, timestamp, result', style='List Bullet')
    doc.add_paragraph('• Data dapat diekspor untuk keperluan pelaporan', style='List Bullet')
    
    doc.add_heading('5.3 Prinsip yang Dijaga', level=2)
    doc.add_paragraph('✅ Tidak ada kredensial terpisah (menggunakan SSO existing)', style='List Bullet')
    doc.add_paragraph('✅ AI tidak membuat keputusan tanpa review manusia', style='List Bullet')
    doc.add_paragraph('✅ Data tracking tidak mengandung informasi sensitif', style='List Bullet')
    doc.add_paragraph('✅ Komponen dapat dinonaktifkan secara independen', style='List Bullet')
    
    # 6. Skenario Keberlanjutan
    doc.add_heading('6. Skenario Keberlanjutan', level=1)
    
    doc.add_heading('6.1 Jika Pilot Berhasil', level=2)
    
    table6 = doc.add_table(rows=5, cols=2)
    add_table_borders(table6)
    
    hdr6 = table6.rows[0].cells
    hdr6[0].text = 'Langkah'
    hdr6[1].text = 'Tindakan'
    for cell in hdr6:
        cell.paragraphs[0].runs[0].bold = True
    
    data6 = [
        ('1', 'Evaluasi bersama Badan Teknologi Data dan Informasi dan unit terkait'),
        ('2', 'Penyusunan roadmap integrasi lebih luas'),
        ('3', 'Transfer ownership ke unit permanen'),
        ('4', 'Penetapan SLA dan standar operasional'),
    ]
    
    for i, (langkah, tindakan) in enumerate(data6):
        row = table6.rows[i + 1].cells
        row[0].text = langkah
        row[1].text = tindakan
    
    doc.add_paragraph()
    
    doc.add_heading('6.2 Jika Pilot Dihentikan', level=2)
    doc.add_paragraph('• Data kursus dapat diekspor', style='List Bullet')
    doc.add_paragraph('• Sertifikat internal tetap valid sebagai catatan', style='List Bullet')
    doc.add_paragraph('• Tidak ada dampak pada sistem existing (LDAP, HR, LMS lain)', style='List Bullet')
    doc.add_paragraph('• Dokumentasi lessons learned disusun', style='List Bullet')
    
    doc.add_heading('6.3 Prinsip Handover ke Badan Teknologi Data dan Informasi', level=2)
    doc.add_paragraph('Jika dilanjutkan dalam skala lebih luas:')
    doc.add_paragraph('• Seluruh aset, dokumentasi, dan source code diserahkan ke Badan Teknologi Data dan Informasi', style='List Bullet')
    doc.add_paragraph('• Transfer knowledge secara terstruktur', style='List Bullet')
    doc.add_paragraph('• Inisiator beralih peran menjadi pendamping', style='List Bullet')
    
    # 7. Penutup
    doc.add_heading('7. Penutup', level=1)
    
    doc.add_paragraph(
        'Dokumen ini menggambarkan kondisi implementasi pilot ULP saat ini dan prinsip-prinsip yang dijaga dalam pengembangannya.'
    )
    
    doc.add_paragraph('Poin-poin utama:', style='Intense Quote')
    doc.add_paragraph('• ULP adalah entry point terpadu, bukan pengganti sistem existing', style='List Bullet')
    doc.add_paragraph('• AI digunakan hanya untuk membantu, bukan mengambil keputusan', style='List Bullet')
    doc.add_paragraph('• Autentikasi menggunakan infrastruktur existing (LDAP)', style='List Bullet')
    doc.add_paragraph('• Sistem dirancang agar dapat dihentikan tanpa merusak sistem lain', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('Kami mengundang seluruh pihak terkait untuk memberikan tanggapan dan saran penyempurnaan.')
    
    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.add_run('Dokumen ini disusun dengan semangat kolaborasi lintas unit untuk kepentingan pengembangan kapabilitas pembelajaran organisasi.').italic = True
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc

# Generate document
doc = create_ulp_document()
output_path = r'C:\Users\yudhiar\.gemini\antigravity\brain\efc2da0d-e0c1-4a3e-a851-b22d6eefbf36\ULP_Concept_Document.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
